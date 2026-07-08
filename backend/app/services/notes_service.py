"""NotesService — orchestrates AI note generation on the shared RAG pipeline.

Responsibilities:
  * resolve source context (semantic search OR direct chunk selection)
  * build prompts via PromptBuilder + PromptRegistry
  * stream the LLM generation (SSE) with cancel/retry support
  * persist notes, versions, generations and cost logs
  * plain CRUD, versioning, restore and export

Retrieval NEVER feeds whole documents to the model — it always goes through
VectorSearchService or ChunkRepository.
"""

from __future__ import annotations

import json
import time
from datetime import datetime, timezone
from typing import AsyncIterator
from uuid import UUID

import structlog
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exceptions import ForbiddenError, NotFoundError, ValidationError
from app.db.session import AsyncSessionLocal
from app.models.note import Note, NoteVersion
from app.models.resource import Resource
from app.models.squad import SquadMember
from app.models.vault import Vault
from app.repositories.chunk_repository import ChunkRepository, ScoredChunk
from app.repositories.note_repository import NoteRepository
from app.schemas.auth import CurrentUser
from app.schemas.note_schema import (
    NoteGenerateRequest,
    NoteUpdateRequest,
)
from app.services import prompt_builder, prompt_registry, vector_search_service
from app.services.ai import get_provider
from app.services.ai.base import TokenUsage
from app.services.ai.pricing import estimate_cost
from app.services.ai.usage_logger import log_generation
from app.services.prompt_builder import GenerationOptions

logger = structlog.get_logger()

NOTE_SEARCH_TOP_K = 30


# ─────────────────────────────────────────────────────────────────────────────
# Authorization helpers
# ─────────────────────────────────────────────────────────────────────────────

async def _assert_vault_member(db: AsyncSession, vault_id: UUID, user_id: UUID) -> Vault:
    result = await db.execute(
        select(Vault).where(Vault.id == vault_id, Vault.deleted_at.is_(None))
    )
    vault = result.scalar_one_or_none()
    if not vault:
        raise NotFoundError("Vault")
    mem = await db.execute(
        select(SquadMember).where(
            SquadMember.squad_id == vault.squad_id,
            SquadMember.user_id == user_id,
            SquadMember.removed_at.is_(None),
        )
    )
    if not mem.scalar_one_or_none():
        raise ForbiddenError("You are not a member of this squad.")
    return vault


async def _get_note_and_authorize(db: AsyncSession, note_id: UUID, user_id: UUID) -> Note:
    repo = NoteRepository(db)
    note = await repo.get(note_id)
    if not note:
        raise NotFoundError("Note")
    await _assert_vault_member(db, note.vault_id, user_id)
    return note


async def _assert_resources_in_vault(
    db: AsyncSession, vault_id: UUID, resource_ids: list[UUID]
) -> None:
    if not resource_ids:
        return
    result = await db.execute(
        select(Resource.id).where(
            Resource.id.in_(resource_ids),
            Resource.vault_id == vault_id,
            Resource.deleted_at.is_(None),
        )
    )
    found = {row[0] for row in result.all()}
    missing = set(resource_ids) - found
    if missing:
        raise ForbiddenError("One or more selected resources do not belong to this vault.")


# ─────────────────────────────────────────────────────────────────────────────
# Retrieval
# ─────────────────────────────────────────────────────────────────────────────

def _wrap(chunks) -> list[ScoredChunk]:
    return [ScoredChunk(chunk=c, similarity=1.0) for c in chunks]


async def _resolve_titles(db: AsyncSession, resource_ids: set[UUID]) -> dict[UUID, str]:
    if not resource_ids:
        return {}
    result = await db.execute(
        select(Resource.id, Resource.title).where(Resource.id.in_(resource_ids))
    )
    return {row[0]: row[1] for row in result.all()}


async def _retrieve(
    db: AsyncSession, vault: Vault, user_id: UUID, req: NoteGenerateRequest
) -> list[ScoredChunk]:
    repo = ChunkRepository(db)
    limit = req.max_context_chunks

    if req.retrieval_mode == "vault":
        query = req.query or "key concepts, definitions and important points"
        return await vector_search_service.search(
            db, vault_id=vault.id, query=query, user_id=user_id,
            top_k=NOTE_SEARCH_TOP_K, rerank_k=limit,
        )

    if req.retrieval_mode == "resources":
        await _assert_resources_in_vault(db, vault.id, req.resource_ids)
        if not req.resource_ids:
            raise ValidationError("Select at least one resource.")
        if req.query:
            return await vector_search_service.search(
                db, vault_id=vault.id, query=req.query, user_id=user_id,
                top_k=NOTE_SEARCH_TOP_K, rerank_k=limit, resource_ids=req.resource_ids,
            )
        return _wrap(await repo.get_by_resources(req.resource_ids, limit=limit))

    if req.retrieval_mode == "chapters":
        await _assert_resources_in_vault(db, vault.id, req.resource_ids)
        if not req.resource_ids or not req.chapters:
            raise ValidationError("Select resources and at least one chapter.")
        return _wrap(await repo.get_by_headings(req.resource_ids, req.chapters, limit=limit))

    if req.retrieval_mode == "pages":
        if not req.page_resource_id or not req.pages:
            raise ValidationError("Select a resource and at least one page.")
        await _assert_resources_in_vault(db, vault.id, [req.page_resource_id])
        return _wrap(await repo.get_by_pages(req.page_resource_id, req.pages, limit=limit))

    raise ValidationError(f"Unknown retrieval_mode: {req.retrieval_mode}")


def _options_from_settings(req: NoteGenerateRequest) -> GenerationOptions:
    s = req.settings
    return GenerationOptions(
        length=s.length,
        difficulty=s.difficulty,
        audience=s.audience,
        language=s.language,
        output_format=s.output_format,
        tone=s.tone,
        exam_focus=s.exam_focus,
        include_citations=s.include_citations,
    )


def _default_title(req: NoteGenerateRequest, template_label: str) -> str:
    if req.title:
        return req.title.strip()
    if req.query:
        return f"{template_label}: {req.query.strip()[:60]}"
    return template_label


# ─────────────────────────────────────────────────────────────────────────────
# Streaming generation (SSE)
# ─────────────────────────────────────────────────────────────────────────────

def _sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"


async def stream_generate(
    user: CurrentUser,
    vault_id: UUID,
    req: NoteGenerateRequest,
    *,
    note_id: UUID | None = None,
) -> AsyncIterator[str]:
    """Generate (or regenerate) a note, streaming SSE events.

    Events: ``meta`` (citations/model) → ``delta``* (text) → ``done`` (note_id +
    usage), or ``error``. Owns its own DB session so the stream is not bound to
    the request-scoped session lifecycle.
    """
    async with AsyncSessionLocal() as db:
        try:
            vault = await _assert_vault_member(db, vault_id, user.id)
            template = prompt_registry.get_template(req.mode)
            options = _options_from_settings(req)

            existing_note: Note | None = None
            if note_id is not None:
                existing_note = await _get_note_and_authorize(db, note_id, user.id)
                if existing_note.vault_id != vault_id:
                    raise ForbiddenError("Note does not belong to this vault.")

            scored = await _retrieve(db, vault, user.id, req)
            if not scored:
                yield _sse("error", {
                    "message": "No AI-ready content found for this selection. "
                               "Make sure the resources have finished processing.",
                })
                return

            titles = await _resolve_titles(db, {sc.chunk.resource_id for sc in scored})
            bundle = prompt_builder.build_prompt_bundle(
                template, options, topic=req.query, scored=scored, resource_titles=titles
            )

            yield _sse("meta", {
                "mode": template.key,
                "model": None,  # filled by provider; informational
                "citations": [_citation_dict(c) for c in bundle.citations],
            })

            provider = get_provider()
            usage = TokenUsage()
            parts: list[str] = []
            started = time.perf_counter()

            async for evt in provider.stream_chat(
                bundle.messages, temperature=0.3,
            ):
                if evt.type == "delta":
                    parts.append(evt.text)
                    yield _sse("delta", {"text": evt.text})
                elif evt.type == "done":
                    usage = evt.usage

            content = "".join(parts).strip()
            latency_ms = int((time.perf_counter() - started) * 1000)
            if not content:
                yield _sse("error", {"message": "The model returned no content."})
                return

            note, gen_row = await _persist_generation(
                db, user=user, vault_id=vault_id, req=req, template=template,
                content=content, usage=usage, latency_ms=latency_ms,
                citations=bundle.citations, existing_note=existing_note,
            )
            await db.commit()

            model = get_provider().name
            from app.core.config import settings as _settings
            yield _sse("done", {
                "note_id": str(note.id),
                "word_count": note.word_count,
                "generation": {
                    "prompt_tokens": usage.prompt_tokens,
                    "completion_tokens": usage.completion_tokens,
                    "total_tokens": usage.total_tokens,
                    "cost_usd": estimate_cost(
                        _settings.LLM_CHAT_MODEL, usage.prompt_tokens, usage.completion_tokens
                    ),
                    "latency_ms": latency_ms,
                    "model": _settings.LLM_CHAT_MODEL,
                    "provider": model,
                },
            })
        except (ForbiddenError, NotFoundError, ValidationError) as e:
            await db.rollback()
            yield _sse("error", {"message": e.detail})
        except Exception as e:  # pragma: no cover - safety net
            await db.rollback()
            logger.error("notes.generate_failed", error=str(e))
            yield _sse("error", {"message": "Generation failed. Please retry."})


def _citation_dict(c: prompt_builder.Citation) -> dict:
    return {
        "index": c.index,
        "chunk_id": str(c.chunk_id),
        "resource_id": str(c.resource_id),
        "resource_title": c.resource_title,
        "page_number": c.page_number,
        "heading": c.heading,
        "snippet": c.snippet,
        "similarity": c.similarity,
    }


async def _persist_generation(
    db: AsyncSession,
    *,
    user: CurrentUser,
    vault_id: UUID,
    req: NoteGenerateRequest,
    template,
    content: str,
    usage: TokenUsage,
    latency_ms: int,
    citations: list[prompt_builder.Citation],
    existing_note: Note | None,
) -> tuple[Note, object]:
    from app.core.config import settings

    repo = NoteRepository(db)
    word_count = len(content.split())
    source_resource_ids = list({c.resource_id for c in citations})

    gen_log = await log_generation(
        db,
        user_id=user.id,
        vault_id=vault_id,
        generation_type="note_generation",
        model=settings.LLM_CHAT_MODEL,
        usage=usage,
        latency_ms=latency_ms,
        metadata={
            "mode": template.key,
            "template_version": template.version,
            "retrieval_mode": req.retrieval_mode,
            "citations": len(citations),
        },
    )

    if existing_note is None:
        note = Note(
            vault_id=vault_id,
            created_by=user.id,
            title=_default_title(req, template.label),
            content=content,
            content_format="markdown",
            source_type="ai_generated",
            word_count=word_count,
            metadata_={
                "mode": template.key,
                "retrieval_mode": req.retrieval_mode,
                "settings": req.settings.model_dump(),
            },
        )
        await repo.create(note)
        change_summary = f"Initial generation ({template.label})"
    else:
        note = existing_note
        note.content = content
        note.word_count = word_count
        note.source_type = "ai_generated" if note.source_type == "manual" else note.source_type
        note.metadata_ = {**(note.metadata_ or {}), "last_mode": template.key}
        change_summary = f"Regenerated ({template.label})"

    await repo.create_version(
        note_id=note.id, content=content, created_by=user.id, change_summary=change_summary
    )
    await repo.create_generation(
        note_id=note.id,
        generation_id=gen_log.id,
        source_resource_ids=source_resource_ids,
        prompt_template=f"{template.key}@{template.version}",
    )
    return note, gen_log


# ─────────────────────────────────────────────────────────────────────────────
# CRUD / versions / export (non-streaming)
# ─────────────────────────────────────────────────────────────────────────────

async def list_notes(
    db: AsyncSession, user: CurrentUser, vault_id: UUID, *,
    search: str | None = None, source_type: str | None = None, pinned_only: bool = False,
) -> list[Note]:
    await _assert_vault_member(db, vault_id, user.id)
    return await NoteRepository(db).list_for_vault(
        vault_id, search=search, source_type=source_type, pinned_only=pinned_only
    )


async def get_note(db: AsyncSession, user: CurrentUser, note_id: UUID) -> Note:
    return await _get_note_and_authorize(db, note_id, user.id)


async def create_manual_note(
    db: AsyncSession, user: CurrentUser, vault_id: UUID, *, title: str, content: str = ""
) -> Note:
    await _assert_vault_member(db, vault_id, user.id)
    repo = NoteRepository(db)
    note = Note(
        vault_id=vault_id,
        created_by=user.id,
        title=title.strip(),
        content=content,
        content_format="markdown",
        source_type="manual",
        word_count=len(content.split()),
    )
    await repo.create(note)
    if content.strip():
        await repo.create_version(
            note_id=note.id, content=content, created_by=user.id, change_summary="Created"
        )
    await db.commit()
    await db.refresh(note)
    return note


async def update_note(
    db: AsyncSession, user: CurrentUser, note_id: UUID, data: NoteUpdateRequest
) -> Note:
    note = await _get_note_and_authorize(db, note_id, user.id)
    repo = NoteRepository(db)

    content_changed = data.content is not None and data.content != note.content
    if content_changed:
        # Snapshot the new content as a version; mark hybrid if it was AI-made.
        note.content = data.content  # type: ignore[assignment]
        note.word_count = len(data.content.split())  # type: ignore[union-attr]
        if note.source_type == "ai_generated":
            note.source_type = "hybrid"
        await repo.create_version(
            note_id=note.id, content=note.content, created_by=user.id,
            change_summary=data.change_summary or "Manual edit",
        )
    if data.title is not None:
        note.title = data.title.strip()
    if data.is_pinned is not None:
        note.is_pinned = data.is_pinned

    await db.commit()
    await db.refresh(note)
    return note


async def delete_note(db: AsyncSession, user: CurrentUser, note_id: UUID) -> None:
    note = await _get_note_and_authorize(db, note_id, user.id)
    await NoteRepository(db).soft_delete(note)
    await db.commit()


async def list_versions(db: AsyncSession, user: CurrentUser, note_id: UUID) -> list[NoteVersion]:
    await _get_note_and_authorize(db, note_id, user.id)
    return await NoteRepository(db).list_versions(note_id)


async def list_generations(db: AsyncSession, user: CurrentUser, note_id: UUID):
    await _get_note_and_authorize(db, note_id, user.id)
    return await NoteRepository(db).list_generations(note_id)


async def restore_version(
    db: AsyncSession, user: CurrentUser, note_id: UUID, version_id: UUID
) -> Note:
    note = await _get_note_and_authorize(db, note_id, user.id)
    repo = NoteRepository(db)
    version = await repo.get_version(version_id)
    if not version or version.note_id != note_id:
        raise NotFoundError("Version")
    note.content = version.content
    note.word_count = len(version.content.split())
    await repo.create_version(
        note_id=note.id, content=version.content, created_by=user.id,
        change_summary=f"Restored v{version.version_number}",
    )
    await db.commit()
    await db.refresh(note)
    return note


async def export_note(
    db: AsyncSession, user: CurrentUser, note_id: UUID, fmt: str
) -> tuple[bytes, str, str]:
    """Return (bytes, filename, content_type) for the requested format."""
    note = await _get_note_and_authorize(db, note_id, user.id)
    safe = "".join(c if c.isalnum() or c in " -_" else "_" for c in note.title).strip() or "note"

    if fmt == "markdown":
        return note.content.encode("utf-8"), f"{safe}.md", "text/markdown"

    if fmt == "docx":
        return _to_docx(note.title, note.content), f"{safe}.docx", (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if fmt == "pdf":
        data = _to_pdf(note.title, note.content)
        if data is None:
            raise ValidationError("PDF export is unavailable (reportlab not installed).")
        return data, f"{safe}.pdf", "application/pdf"

    raise ValidationError(f"Unsupported export format: {fmt}")


def _to_docx(title: str, content: str) -> bytes:
    import io

    import docx

    doc = docx.Document()
    doc.add_heading(title, level=0)
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _to_pdf(title: str, content: str) -> bytes | None:
    try:
        import io

        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception:
        return None

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for para in content.split("\n\n"):
        text = para.strip().replace("\n", "<br/>")
        if text:
            story.append(Paragraph(text, styles["BodyText"]))
            story.append(Spacer(1, 8))
    doc.build(story)
    return buf.getvalue()
